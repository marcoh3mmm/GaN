from docx import Document
from docx.shared import Inches
import os

# Open the MS Word files created in the translations
def open_word_doc(file_name):
    document = Document(file_name)
    return document

# Get the latitudes according to the images names in the GaN webpage (north ="10", south ="10s")
def transform_latitude(lat):
    if "N" in lat:
        lat = str(lat.rstrip(lat[-1]))
    else:
        lat = str(lat.lower())
    return lat


# get the links from the charts for each latitude
def open_local_image(file_name):
    constellation = file_name.split('_')[-4].lower().replace(" ", "-")
    latitude = file_name.split('_')[-2]

    path = os.getcwd() 
    path = os.path.join(path + "\images_local")

    lat = transform_latitude(latitude)

    magnitudes = ["05", "15", "25", "35", "45", "55", "65", "75"]

    paths_list = []
    for mag in magnitudes:
        paths_list.append(path + "\\" + constellation + "-" + lat + "_" + mag + ".png")
    
    return paths_list


#crops an image based on the magnitude, saves image as png
def print_local_image(file_name):
    
    constellation = file_name.split('_')[-4]
    latitude = file_name.split('_')[-2]

    working_doc = open_word_doc(file_name)
    if "Crux_lat_0_" in file_name:
        charts_dir = open_local_image(file_name.replace("0_", "10S_"))
    elif "Bootes_lat_40S_" in file_name:
        charts_dir = open_local_image(file_name.replace("40S_", "30S_"))
    elif "Hercules_lat_40S_" in file_name:
        charts_dir = open_local_image(file_name.replace("40S_", "30S_"))
    else:
        charts_dir = open_local_image(file_name)

    table1 = working_doc.tables[0]
    table1 = (table1.cell(1,0), table1.cell(1,2), table1.cell(4,0), table1.cell(4,2))
    i= 0
    for table_cell in table1:
        table_cell.paragraphs[1].clear()
        table_cell.paragraphs[1].add_run().add_picture(charts_dir[i], width = Inches(3.39), height = Inches(2.35))
        i = i + 1
    working_doc.save(file_name)

    table2 = working_doc.tables[1]
    table2 = (table2.cell(1,0), table2.cell(1,2), table2.cell(4,0), table2.cell(4,2))
    for table_cell in table2:
        table_cell.paragraphs[1].clear()
        table_cell.paragraphs[1].add_run().add_picture(charts_dir[i], width = Inches(3.39), height = Inches(2.35))
        i = i + 1
    working_doc.save(file_name)

    table3 = working_doc.tables[2]
    table3 = (table3.cell(1,0), table3.cell(1,1), table3.cell(1,2), table3.cell(1,3), table3.cell(3,0), table3.cell(3,1), table3.cell(3,2), table3.cell(3,3))
    j = 0
    for table_cell in table3:
        table_cell.paragraphs[0].clear()
        table_cell.paragraphs[0].add_run().add_picture(charts_dir[j], width = Inches(1.44), height = Inches(1.01))
        j = j + 1
    working_doc.save(file_name)


    print("The charts in the activity guide for the constellation {cons}".format(cons = constellation) + " in the latitude {lat}".format(lat = latitude) +" have been printed \n________________________________________________________________________________________________\n")    

    return file_name

    # get the links from the charts for each latitude
def open_download_image(file_name):
    constellation = file_name.split('_')[-4].lower().replace(" ", "-")
    latitude = file_name.split('_')[-2]

    path = os.getcwd() 
    path = os.path.join(path + "\images")

    lat = transform_latitude(latitude)

    magnitudes = ["05", "15", "25", "35", "45", "55", "65", "75"]

    paths_list = []
    for mag in magnitudes:
        paths_list.append(path + "\\" + constellation + "-" + lat + "_" + mag + ".png")
    
    return paths_list


# drops an image based on the magnitude, saves image as png
def print_download_image(file_name):
    
    constellation = file_name.split('_')[-4]
    latitude = file_name.split('_')[-2]

    working_doc = open_word_doc(file_name)
    if "Crux_lat_0_" in file_name:
        charts_dir = open_local_image(file_name.replace("0_", "10S_"))
    elif "Bootes_lat_40S_" in file_name:
        charts_dir = open_local_image(file_name.replace("40S_", "30S_"))
    elif "Hercules_lat_40S_" in file_name:
        charts_dir = open_local_image(file_name.replace("40S_", "30S_"))
    else:
        charts_dir = open_download_image(file_name)

    # working on the tables in the files
    table1 = working_doc.tables[0]
    table1 = (table1.cell(1,0), table1.cell(1,2), table1.cell(4,0), table1.cell(4,2))
    i= 0
    for table_cell in table1:
        table_cell.paragraphs[1].clear()
        table_cell.paragraphs[1].add_run().add_picture(charts_dir[i], width = Inches(3.39), height = Inches(2.35))
        i = i + 1
    working_doc.save(file_name)

    table2 = working_doc.tables[1]
    table2 = (table2.cell(1,0), table2.cell(1,2), table2.cell(4,0), table2.cell(4,2))
    for table_cell in table2:
        table_cell.paragraphs[1].clear()
        table_cell.paragraphs[1].add_run().add_picture(charts_dir[i], width = Inches(3.39), height = Inches(2.35))
        i = i + 1
    working_doc.save(file_name)

    table3 = working_doc.tables[2]
    table3 = (table3.cell(1,0), table3.cell(1,1), table3.cell(1,2), table3.cell(1,3), table3.cell(3,0), table3.cell(3,1), table3.cell(3,2), table3.cell(3,3))
    j = 0
    for table_cell in table3:
        table_cell.paragraphs[0].clear()
        table_cell.paragraphs[0].add_run().add_picture(charts_dir[j], width = Inches(1.44), height = Inches(1.01))
        j = j + 1
    working_doc.save(file_name)


    print("The charts in the activity guide for the constellation {cons}".format(cons = constellation) + " in the latitude {lat}".format(lat = latitude) +" have been printed \n________________________________________________________________________________________________\n")    

    return file_name




