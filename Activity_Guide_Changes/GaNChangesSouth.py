# -*- coding:utf-8 -*-
# Install mtranslate, googletrans for translations
# Install python-docx for managing the Word Files.
# Install Pandas to manage the Excel file and bring the information
# Import Shutil to remove the directory

import os
from deep_translator import GoogleTranslator  
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
import pandas as pd
from shutil import rmtree


def import_south_data():
    # Define the path for the excel file
    excel_path = os.getcwd()
    excel_path = os.path.join(excel_path + "\GaN\Activity_Guide_Changes\GaN_cons_and_dates.xlsx")

    # Get Data from the Excel File using Pandas
    # Capitalize  constellations names for a later comparison
    df_south = pd.read_excel (excel_path,sheet_name='South')
    df_south['Constellations'] = df_south['Constellations'].str.capitalize()

    #store constellation and date information in respective variables
    south_cons = df_south['Constellations']
    south_dates = df_south['Dates']

    #updates the constellation: creates a variable to hold the new constellation, uses the old constellation
    #information to find the new constellation
    new_constellation_south = {}
    
    # Ading key and values to the new dictionary
    for i in range(len(south_cons)):
        new_constellation_south[south_cons[i]] = south_dates[i] 

    #return the dictionary with the North Data
    return new_constellation_south


# Create a list with the paths that will be use by the word files
def create_south_paths(directories, languages, latitudes):

    dir_paths = []
    for lang in languages:
        for direc in directories:
            for lat in latitudes:
                dir_paths.append(direc + "_" + lat + "_" + lang)
    return dir_paths


#opens document that will be edited
def open_word_doc2(file_name):
    document = Document(file_name)
    return document


#updating southern hemisphere information (constellation, date, text displayed to user)
#######################################################################################
########################All the information that needs to change#######################
#######################################################################################

def south_translations(dir_paths):
    #updating southern hemisphere information (constellation, date, text displayed to user)
    south_constellation_replacement = {
        "ChileanSpanish" : "Escorpio",
        "English" : "Scorpius",
        "French" : "Scorpion",
        "Indonesian" : "Scorpio",
        "Portuguese" : "Escorpião",
        "Spanish" : "Escorpio"
        }
        
    south_date_replacement = {
        "ChileanSpanish" : "Del 4 al 13 de julio y del 2 al 11 de agosto",
        "English" : "July 4-13 and August 2-11",
        "French" : "Du 4 au 13 juillet et du 2 au 11 août",
        "Indonesian" : "4 Juli - 13 Juli dan 2 Agustus - 11 Agustus",
        "Portuguese" : "4 a 13 de julho e 2 a 11 de agosto.",
        "Spanish" : "del 4 al 13 de julio y del 2 al 11 de agosto"
        }

    south_heading_first = {
        "ChileanSpanish" : "",
        "English" : " Campaign Dates that use the ",
        "French" : "Dates de la campagne ",
        "Indonesian" : "Waktu Kampanye ",
        "Portuguese" : "Datas das campanhas de ",
        "Spanish" : "Fechas de la campaña año "
        }

    south_heading_middle = {
        "ChileanSpanish" : " Fechas de campaña para la constelación del ",
        "English" : "",
        "French" : " qui utilisent la ",
        "Indonesian" : " yang menggunakan ",
        "Portuguese" : " que usam a ",
        "Spanish" : " que utilizan la "
        }

    south_heading_last = {
        "ChileanSpanish" : ": ",
        "English" : ": ",
        "French" : ": ",
        "Indonesian" : ": ",
        "Portuguese" : ": ",
        "Spanish" : ": "
        }

    paragraph_first = {
        "ChileanSpanish" : "Usted está participando en una campaña mundial para observar y registrar las estrellas visibles más débiles como un medio para medir la contaminación lumínica en un lugar determinado. Localizando y observando la ",
        "English" : "You are participating in a global campaign to observe and record the faintest stars visible as a means of measuring light pollution in a given location. By locating and observing the constellation ",
        "French" :   "Vous allez participer à une campagne mondiale d’observation pour détecter les plus faibles étoiles visibles afin de mesurer la pollution lumineuse sur un site donné. Partout dans le monde, en localisant et en observant la ",
        "Indonesian" : "Anda sedang berpartisipasi dalam kampanye global pengamatan dan pencatatan penampakan bintang paling redup untuk pengukuran tingkat polusi cahaya di suatu lokasi. Melalui pengamatan dan identifikasi  ",
        "Portuguese" : "Está a participar numa campanha global para observar e registar as estrelas mais fracas visíveis como forma de medir a poluição luminosa num determinado local. Localizando e observando a  ",
        "Spanish" : "Usted está participando en una campaña mundial para observar y registrar las estrellas visibles más débiles como un medio para medir la contaminación lumínica en un lugar determinado. Localizando y observando la  ",

    }

    paragraph_last = {
        "ChileanSpanish" : " en el cielo nocturno y comparándolo con las cartas estelares, la gente de todo el mundo aprenderá cómo las luces de su comunidad contribuyen a la contaminación lumínica. Sus contribuciones a la base de datos en línea documentarán el cielo nocturno visible.",
        "English" : " in the night sky and comparing it to stellar charts, people from around the world will learn how the lights in their community contribute to light pollution. Your contributions to the online database will document the visible nighttime sky.",
        "French" :   " dans le ciel nocturne et en la comparant aux cartes stellaires, les participants, apprendront comment l’éclairage, dans leur environnement local, influence la pollution lumineuse. Vos contributions à la base de données en ligne permettront de mesurer la qualité du ciel nocturne.",
        "Indonesian" : " di langit malam dan membandingkannya dengan peta bintang, masyarakat di seluruh dunia dapat mengetahui dan mempelajari seberapa besar kontribusi cahaya di lingkungannya terhadap polusi cahaya. Kontribusi data anda pada basis data online akan membantu mendokumentasikan langit malam yang tampak di berbagai lokasi.",
        "Portuguese" : " no céu noturno e,  comparando-a com cartas estelares, pessoas de todo o mundo aprenderão  como as luzes da sua comunidade contribuem para a poluição luminosa. As suas contribuições para a base de dados on-line irão documentar a visibilidade do céu noturno em todo o mundo.",
        "Spanish" : " en el cielo nocturno y comparándolo con las cartas estelares, la gente de todo el mundo aprenderán cómo las luces de su comunidad contribuyen a la contaminación lumínica. Sus contribuciones a la base de datos en línea documentarán el cielo nocturno visible.",

    }

    ##################################################################################################
    ##################################################################################################
            ###	End of the changes section defining things that need to be changed			###
    ##################################################################################################
    ##################################################################################################

    # Get data from the Excel file and bring the created Paths
    dir_path = dir_paths
    South_data = import_south_data()


# Organize the Languages by lists to make better translations
    country_list1 = ("Chilean_Spanish", "French", "Indonesian", "Portuguese","Spanish")
    country_list2 = ("English")
    

    # Getting data from the Paths
    language_base = dir_path.split('_')[-1]
    latitude = dir_path.split('_')[-2]
    const_name = dir_path.split('_')[-3]
    year = dir_path.split('_')[-5]
    #thaiYear = int(year)+ 543

    #Be sure to change the websites into the word files
    website1 = "astro/maps/GaNight/2018/"
    website2 = "astro/maps/GaNight/2019/"

    
    # Define the Word file path as the original file
    word_path = os.path.abspath("..\Gan\GaN\docs_to_change\GaN2018_ActivityGuide_Scorpius_S_")
    working_doc = open_word_doc2(word_path + str(language_base) + ".docx")

    # styles of each paragraph to kkep the original word styles
    obj_styles = working_doc.styles
    obj_char_style = obj_styles.add_style('GaN_style', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_char_style.font
    obj_font.name = 'Calibri'
    obj_font.size = Pt(11)
    
    obj_styles2 = working_doc.styles
    obj_char_style2 = obj_styles2.add_style('GaN_paragraph', WD_STYLE_TYPE.CHARACTER)
    obj_font2 = obj_char_style2.font
    obj_font2.name = 'Calibri'
    obj_font2.size = Pt(10)
    
    obj_styles3 = working_doc.styles
    obj_char_style3 = obj_styles3.add_style('GaN_links', WD_STYLE_TYPE.CHARACTER)
    obj_font3 = obj_char_style3.font
    obj_font3.name = 'Calibri'
    obj_font3.size = Pt(9.5)
    obj_font3.bold = True
    obj_font3.underline = True
    obj_font3.color.rgb = RGBColor(0,0,128)


    #Define the base language in deep_translator and translate it into de destiny language
    if const_name != "Canis Major" :
        constellation_translated =GoogleTranslator(source ='english', target = language_base.lower()).translate(const_name +" constellation")
        date_translated = GoogleTranslator(source ='english', target = language_base.lower()).translate(South_data.get(const_name))
    else:
        constellation_translated = "Canis Major"
        date_translated = South_data.get('Canis major')


    # Replace the translations in the proper places
    for language_selected, date in south_date_replacement.items():
        if language_selected == language_base:
            for paragraph in working_doc.paragraphs:
                #If the contellation's name is in the paragraph, delete the paragraph and add a new one with the translations
                if south_constellation_replacement[language_base] in paragraph.text:
                    # Replace only if the name and the date is on the paragraph, organizng with the grammar of each language
                    if date in paragraph.text:
                        paragraph.clear()
                        if language_base in country_list1:
                            paragraph.add_run(south_heading_first[language_base]+ str(year) + south_heading_middle[language_base]+ constellation_translated + south_heading_last[language_base] + date_translated + ".", style = 'GaN_style')
                        elif language_base in country_list2:
                            paragraph.add_run(str(year) + south_heading_first[language_base]+ constellation_translated + south_heading_middle[language_base] + south_heading_last[language_base]+ date_translated + ".", style = 'GaN_style' )
                    
                    # Replace only if the constellation's name is in the paragraph
                    else:
                        paragraph.clear()
                        if(language_base!= 'Japanese'):
                            paragraph.add_run(paragraph_first[language_base] + constellation_translated + paragraph_last[language_base], style = 'GaN_paragraph')
                        else:
                            paragraph.add_run(paragraph_first[language_base] + paragraph_last[language_base] + constellation_translated, style = 'GaN_paragraph')
                
                if website1 in paragraph.text:
                    new_link = paragraph.text.replace("2018",str(year))
                    paragraph.text = None
                    paragraph.add_run(new_link, style = 'GaN_links')
                
                elif website2 in paragraph.text:
                    new_link = paragraph.text.replace("2019",str(year))
                    paragraph.text = None
                    paragraph.add_run(new_link, style = 'GaN_links')

    #Save a copy with a new name, date and language.
    dir_path = dir_path.rsplit('_', 2)[0]
    new_word_path = os.path.join(dir_path + "\\GaN_{year}_ActivityGuide_{cons}_lat_".format(year = year, cons = const_name) + str(latitude) + "_" + str(language_base) + ".docx")
    working_doc.save(new_word_path)


    #Print information about the working file on
    print("The " + language_base + " activity guide for the constellation {cons}".format(cons = const_name) + " in the latitude {lat}".format(lat = latitude) +" south has been completed \n___________________________________________________________________________________________________________\n")

    # return the new doc path to make a list with it.
    return new_word_path



